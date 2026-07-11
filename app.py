from flask import Flask, render_template, request, redirect, session, flash, url_for, Response, jsonify
from config import Config
from datetime import datetime
from datetime import date
from io import BytesIO, StringIO
from types import SimpleNamespace

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from sqlalchemy import func
from sqlalchemy import extract
from models.user import db, User
from models.pengiriman import Pengiriman
from models.transaksi import Barang, BarangMasuk, Transaksi, DetailTransaksi, Retur
from models.kritik import Kritik

app = Flask(__name__)


from config import Config
app.config.from_object(Config)

db.init_app(app)
@app.teardown_appcontext
def shutdown_session(exception=None):
    db.session.remove()

def init_default_data():
    with app.app_context():
        db.create_all()
        if not User.query.filter_by(username="admin").first():
            db.session.add(User(username="admin", password="admin", role="admin"))
            db.session.commit()


init_default_data()


@app.route("/")
def index():
    return redirect(url_for("login"))


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()

        user = User.query.filter_by(username=username, password=password).first()

        if user:
            session["user_id"] = user.id
            session["username"] = user.username
            session["role"] = user.role
             

            role = user.role.lower()
            if role == "admin":
                return redirect("/admin/dashboard")
            elif role == "kasir":
                return redirect("/kasir/dashboard")
            elif role in {"staf gudang", "gudang"}:
                return redirect("/gudang/dashboard")

        return render_template("login.html", error="Username atau Password Salah")

    return render_template("login.html")


@app.route("/admin/dashboard")
def admin_dashboard():
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() != "admin":
        return "Akses Ditolak: Anda bukan Admin", 403

    total_barang = Barang.query.count()
    total_stok = Barang.query.with_entities(db.func.sum(Barang.stok)).scalar() or 0
    total_nilai_stok = sum(item.harga * item.stok for item in Barang.query.all())
    stok_menipis = Barang.query.filter(Barang.stok <= 5).order_by(Barang.stok).all()
    barang_terbaru = Barang.query.order_by(Barang.id.desc()).limit(5).all()

    return render_template(
        "admin/dashboard.html",
        total_barang=total_barang,
        total_stok=total_stok,
        total_nilai_stok=total_nilai_stok,
        stok_menipis=stok_menipis,
        barang_terbaru=barang_terbaru,
    )


@app.route("/admin/barang", methods=["GET", "POST"])
def admin_barang():
    if request.method == "POST":
        nama = request.form.get("nama_barang", "").strip()
        harga = request.form.get("harga", "0")
        stok = request.form.get("stok", "0")
        satuan = request.form.get("satuan", "Pcs")
        tgl_masuk = request.form.get("tanggal_masuk")
        tgl_expired = request.form.get("tanggal_kadaluarsa")
        

        penginput = session.get("user_id")

        if nama:
            barang_baru = Barang(
                nama_barang=nama,
                harga=int(harga or 0),
                stok=int(stok or 0),
                satuan=satuan,
                tanggal_masuk=tgl_masuk or "-",
                tanggal_kadaluarsa=tgl_expired or "-",
                id_admin_fk=penginput,
            )
            db.session.add(barang_baru)
            db.session.commit()
        return redirect(url_for("admin_barang"))

    search_query = request.args.get("q", "").strip()
    query = Barang.query
    if search_query:
        query = query.filter(Barang.nama_barang.ilike(f"%{search_query}%"))

    barang_list = query.order_by(Barang.id.desc()).all()
    return render_template("admin/barang.html", barang_list=barang_list, search_query=search_query)


@app.route("/admin/barang/hapus/<int:item_id>")
def hapus_barang(item_id):
    barang = Barang.query.get_or_404(item_id)
    db.session.delete(barang)
    db.session.commit()
    return redirect(url_for("admin_barang"))


@app.route("/admin/barang/edit/<int:item_id>", methods=["GET", "POST"])
def edit_barang(item_id):
    barang = Barang.query.get_or_404(item_id)

    if request.method == "POST":
        barang.nama_barang = request.form.get("nama_barang", barang.nama_barang).strip()
        barang.harga = int(request.form.get("harga", barang.harga) or barang.harga)
        barang.stok = int(request.form.get("stok", barang.stok) or barang.stok)
        barang.satuan = request.form.get("satuan", barang.satuan)
        barang.tanggal_masuk = request.form.get("tanggal_masuk", barang.tanggal_masuk)
        barang.tanggal_kadaluarsa = request.form.get("tanggal_kadaluarsa", barang.tanggal_kadaluarsa)
        db.session.commit()
        return redirect(url_for("admin_barang"))

    return render_template("admin/edit_barang.html", barang=barang)


@app.route("/admin/user", methods=["GET", "POST"])
def admin_user():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "").strip()
        role = request.form.get("role", "").strip()

        if username and password and role:
            existing_user = User.query.filter_by(username=username).first()
            if not existing_user:
                new_user = User(username=username, password=password, role=role)
                db.session.add(new_user)
                db.session.commit()

        return redirect(url_for("admin_user"))

    search_query = request.args.get("q", "").strip()
    query = User.query
    if search_query:
        query = query.filter(
            User.username.ilike(f"%{search_query}%")
            | User.role.ilike(f"%{search_query}%")
        )

    users = query.order_by(User.id).all()
    return render_template("admin/user.html", users=users, search_query=search_query)
@app.route("/admin/user/hapus/<int:item_id>")
def hapus_user(item_id):
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() != "admin":
        return "Akses Ditolak: Anda bukan Admin", 403

    user = User.query.get_or_404(item_id)

    if user.id == session.get("user_id"):
        flash("Tidak bisa menghapus akun yang sedang login.", "error")
        return redirect(url_for("admin_user"))

    db.session.delete(user)
    db.session.commit()
    return redirect(url_for("admin_user"))

@app.route("/admin/laporan")
def admin_laporan():
    total_barang = Barang.query.count()
    total_stok = Barang.query.with_entities(db.func.sum(Barang.stok)).scalar() or 0
    total_nilai = sum(item.harga * item.stok for item in Barang.query.all())
    return render_template(
        "admin/laporan.html",
        total_barang=total_barang,
        total_stok=total_stok,
        total_nilai=total_nilai,
        barang_list=Barang.query.order_by(Barang.id.desc()).all(),
    )


@app.route("/admin/laporan/download")
def admin_laporan_download():
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() != "admin":
        return "Akses Ditolak: Anda bukan Admin", 403

    barang_list = Barang.query.order_by(Barang.id).all()
    si = StringIO()
    writer = csv.writer(si)
    writer.writerow(["ID", "Nama Barang", "Harga", "Stok", "Satuan", "Tanggal Masuk", "Tanggal Kadaluarsa", "Diinput Oleh"])

    for barang in barang_list:
        writer.writerow([
            barang.id,
            barang.nama_barang,
            barang.harga,
            barang.stok,
            barang.satuan,
            barang.tanggal_masuk,
            barang.tanggal_kadaluarsa,
            barang.id_admin_fk,
        ])

    output = si.getvalue()
    return Response(
        output,
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=laporan_barang.csv"},
    )


@app.route("/admin/laporan/download/pdf")
def admin_laporan_download_pdf():
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() != "admin":
        return "Akses Ditolak: Anda bukan Admin", 403

    barang_list = Barang.query.order_by(Barang.id).all()
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph("Laporan Stok Barang", styles["Title"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"Total barang: {len(barang_list)}", styles["BodyText"]))
    story.append(Spacer(1, 12))

    data = [["ID", "Nama Barang", "Harga", "Stok", "Satuan"]]
    for barang in barang_list:
        data.append([str(barang.id), barang.nama_barang, f"Rp {barang.harga:,}", str(barang.stok), barang.satuan])

    table = Table(data, repeatRows=1)
    table.setStyle(
        TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#FFA726")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("PADDING", (0, 0), (-1, -1), 6),
        ])
    )
    story.append(table)

    doc.build(story)
    buffer.seek(0)
    return Response(buffer.getvalue(), mimetype="application/pdf", headers={"Content-Disposition": "attachment; filename=laporan_barang.pdf"})


@app.route("/admin/laporan/download/docx")
def admin_laporan_download_docx():
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() != "admin":
        return "Akses Ditolak: Anda bukan Admin", 403

    barang_list = Barang.query.order_by(Barang.id).all()
    document = Document()
    document.add_heading("Laporan Stok Barang", level=1)
    document.add_paragraph(f"Total barang: {len(barang_list)}")

    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "ID"
    headers[1].text = "Nama Barang"
    headers[2].text = "Harga"
    headers[3].text = "Stok"
    headers[4].text = "Satuan"

    for barang in barang_list:
        row = table.add_row().cells
        row[0].text = str(barang.id)
        row[1].text = barang.nama_barang
        row[2].text = f"Rp {barang.harga:,}"
        row[3].text = str(barang.stok)
        row[4].text = barang.satuan

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return Response(buffer.getvalue(), mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=laporan_barang.docx"})


@app.route("/admin/kritik")
def admin_kritik():
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() != "admin":
        return "Akses Ditolak: Anda bukan Admin", 403

    kritik_list = Kritik.query.order_by(Kritik.id.desc()).all()
    return render_template("admin/kritik.html", kritik_list=kritik_list)


@app.route("/admin/kritik/simpan", methods=["POST"])
def admin_kritik_simpan():
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() != "admin":
        return "Akses Ditolak: Anda bukan Admin", 403

    tanggal = (request.form.get("tanggal") or "").strip()
    nama_pelanggan = (request.form.get("nama_pelanggan") or "").strip()
    feedback = (request.form.get("feedback") or "").strip()
    status = (request.form.get("status") or "Belum Diproses").strip() or "Belum Diproses"

    if not tanggal or not nama_pelanggan or not feedback:
        flash("Tanggal, Nama Pelanggan, dan Feedback wajib diisi.", "error")
        return redirect(url_for("admin_kritik"))

    item = Kritik(
        tanggal=tanggal,
        nama_pelanggan=nama_pelanggan,
        feedback=feedback,
        status=status,
    )
    db.session.add(item)
    db.session.commit()
    return redirect(url_for("admin_kritik"))


@app.route("/admin/kritik/edit/<int:item_id>", methods=["POST"])
def admin_kritik_edit(item_id):
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() != "admin":
        return "Akses Ditolak: Anda bukan Admin", 403

    kritik = Kritik.query.get_or_404(item_id)

    kritik.tanggal = (request.form.get("tanggal") or kritik.tanggal).strip()
    kritik.nama_pelanggan = (request.form.get("nama_pelanggan") or kritik.nama_pelanggan).strip()
    kritik.feedback = (request.form.get("feedback") or kritik.feedback).strip()
    kritik.status = (request.form.get("status") or kritik.status).strip() or kritik.status

    db.session.commit()
    return redirect(url_for("admin_kritik"))


@app.route("/admin/kritik/hapus/<int:item_id>", methods=["POST"])
def admin_kritik_hapus(item_id):
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() != "admin":
        return "Akses Ditolak: Anda bukan Admin", 403

    kritik = Kritik.query.get_or_404(item_id)
    db.session.delete(kritik)
    db.session.commit()
    return redirect(url_for("admin_kritik"))



@app.route("/kasir/dashboard")
def kasir_dashboard():
    if "user_id" not in session:
        return redirect("/login")
    if session.get("role").lower() != "kasir":
        return "Akses Ditolak: Anda bukan Kasir", 403

    hari_ini = date.today()
    transaksi_list = Transaksi.query.filter(
        db.func.date(Transaksi.tanggal) == hari_ini
    ).order_by(Transaksi.tanggal.desc()).all()

    total_penjualan = sum(t.total for t in transaksi_list)
    jumlah_barang = Barang.query.filter(Barang.stok > 0).count()

    return render_template("kasir/dashboard.html",
        transaksi_hari_ini=len(transaksi_list),
        total_penjualan=total_penjualan,
        jumlah_barang=jumlah_barang,
        transaksi_list=transaksi_list
    )


@app.route("/kasir/transaksi", methods=["GET", "POST"])
def kasir_transaksi():
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() != "kasir":
        return "Akses Ditolak: Anda bukan Kasir", 403

    if request.method == "POST":
        payload = request.get_json(silent=True) or {}
        items = payload.get("items", [])
        kasir_id = payload.get("kasir_id")

        if not items:
            return {"error": "Keranjang kosong"}, 400

        total = 0
        for item in items:
            barang = Barang.query.get(item.get("barang_id"))
            if not barang:
                return {"error": "Barang tidak ditemukan"}, 400
            qty = int(item.get("qty", 0))
            if qty <= 0 or qty > barang.stok:
                return {"error": f"Stok {barang.nama_barang} tidak mencukupi"}, 400
            total += barang.harga * qty
            barang.stok -= qty
            

        transaksi_baru = Transaksi(
            kasir_id=session["user_id"],
            total=total,
            tanggal=datetime.now()
        )
        db.session.add(transaksi_baru)
        db.session.flush()

        for item in items:
            barang = db.session.get(Barang, item.get("barang_id"))
            detail = DetailTransaksi(
                transaksi_id=transaksi_baru.id,
                barang_id=item.get("barang_id"),
                qty=int(item.get("qty", 0)),
                subtotal=barang.harga * int(item.get("qty", 0))
            )
            db.session.add(detail)

        db.session.commit()
        return {"ok": True, "total": total}

    barang_list = Barang.query.order_by(Barang.nama_barang).all()
    return render_template("kasir/transaksi.html", barang_list=barang_list)


@app.route("/kasir/retur", methods=["GET", "POST"])
def kasir_retur():
    if "user_id" not in session:
        return redirect("/login")
    if (session.get("role") or "").lower() != "kasir":
        return "Akses Ditolak: Anda bukan Kasir", 403

    if request.method == "POST":
        data = request.get_json(silent=True) or {}
        id_retur = data.get("id_retur", "").strip()
        tanggal = data.get("tanggal", "").strip()
        nama_barang = data.get("nama_barang", "").strip()
        jumlah = data.get("jumlah", 0)
        alasan = data.get("alasan", "").strip()
        status = data.get("status", "Pending").strip()

        if not id_retur or not nama_barang:
            return jsonify({"error": "ID Retur dan Nama Barang wajib diisi!"}), 400

        existing = Retur.query.filter_by(id_retur=id_retur).first()
        if existing:
            return jsonify({"error": f"ID Retur {id_retur} sudah ada!"}), 400

        retur_baru = Retur(
            id_retur=id_retur,
            tanggal=tanggal,
            nama_barang=nama_barang,
            jumlah=int(jumlah) if jumlah else 0,
            alasan=alasan,
            status=status,
            kasir_id=session["user_id"]
        )
        db.session.add(retur_baru)
        db.session.commit()
        return jsonify({"success": True})

    retur_list = Retur.query.order_by(Retur.tanggal.desc()).all()
    return render_template("kasir/retur.html", retur_list=retur_list)

@app.route("/kasir/retur/update-status", methods=["POST"])
def update_status_retur():
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() != "kasir":
        return {"error": "Akses Ditolak"}, 403

    payload = request.get_json(silent=True) or {}
    id_retur = str(payload.get("id_retur", "")).strip()
    status = str(payload.get("status", "Pending")).strip() or "Pending"

    if not id_retur:
        return {"error": "ID retur kosong"}, 400

    retur = Retur.query.filter_by(id_retur=id_retur).first()
    if not retur:
        retur = Retur(
            id_retur=id_retur,
            tanggal=datetime.now().strftime("%d-%m-%Y"),
            nama_barang="",
            jumlah=0,
            alasan="",
            status=status,
            kasir_id=session["user_id"],
        )
        db.session.add(retur)
    else:
        retur.status = status

    db.session.commit()
    return {"success": True}

@app.route("/kasir/pengiriman", methods=["GET", "POST"])
def kasir_pengiriman():
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() != "kasir":
        return "Akses Ditolak: Anda bukan Kasir", 403

    if request.method == "POST":
        no_pengiriman = request.form.get("no_pengiriman", "").strip()
        tanggal_input = request.form.get("tanggal_input", "").strip()
        nama_penerima = request.form.get("nama_penerima", "").strip()
        barang = request.form.get("barang", "").strip()
        jumlah = request.form.get("jumlah", "0").strip()
        status = request.form.get("status", "Diproses").strip()

        if no_pengiriman and nama_penerima and barang:
            existing = Pengiriman.query.filter_by(no_pengiriman=no_pengiriman).first()
            if not existing:
                if tanggal_input:
                    try:
                        tanggal_display = datetime.strptime(tanggal_input, "%Y-%m-%d").strftime("%d-%m-%Y")
                    except ValueError:
                        tanggal_display = tanggal_input
                else:
                    tanggal_display = datetime.now().strftime("%d-%m-%Y")

                pengiriman = Pengiriman(
                    no_pengiriman=no_pengiriman,
                    tanggal_input=tanggal_display,
                    nama_penerima=nama_penerima,
                    barang=barang,
                    jumlah=int(jumlah or 0),
                    status=status or "Diproses",
                )
                db.session.add(pengiriman)
                db.session.commit()

        return redirect(url_for("kasir_pengiriman"))

    pengiriman_list = Pengiriman.query.order_by(Pengiriman.id.desc()).all()
    return render_template("kasir/pengiriman.html", pengiriman_list=pengiriman_list)


@app.route("/kasir/pengiriman/edit/<int:item_id>", methods=["GET", "POST"])
def edit_pengiriman(item_id):
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() != "kasir":
        return "Akses Ditolak: Anda bukan Kasir", 403

    pengiriman = Pengiriman.query.get_or_404(item_id)

    if request.method == "POST":
        pengiriman.no_pengiriman = request.form.get("no_pengiriman", pengiriman.no_pengiriman).strip()
        tanggal_input = request.form.get("tanggal_input", "").strip()
        if tanggal_input:
            try:
                pengiriman.tanggal_input = datetime.strptime(tanggal_input, "%Y-%m-%d").strftime("%d-%m-%Y")
            except ValueError:
                pengiriman.tanggal_input = tanggal_input
        pengiriman.nama_penerima = request.form.get("nama_penerima", pengiriman.nama_penerima).strip()
        pengiriman.barang = request.form.get("barang", pengiriman.barang).strip()
        pengiriman.jumlah = int(request.form.get("jumlah", pengiriman.jumlah) or pengiriman.jumlah)
        pengiriman.status = request.form.get("status", pengiriman.status).strip() or pengiriman.status
        db.session.commit()
        return redirect(url_for("kasir_pengiriman"))

    return render_template("kasir/edit_pengiriman.html", pengiriman=pengiriman)


@app.route("/kasir/pengiriman/hapus/<int:item_id>")
def hapus_pengiriman(item_id):
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() != "kasir":
        return "Akses Ditolak: Anda bukan Kasir", 403

    pengiriman = Pengiriman.query.get_or_404(item_id)
    db.session.delete(pengiriman)
    db.session.commit()
    return redirect(url_for("kasir_pengiriman"))


@app.route("/kasir/pengiriman/status/<int:item_id>", methods=["POST"])
def update_status_pengiriman(item_id):
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() != "kasir":
        return "Akses Ditolak: Anda bukan Kasir", 403

    pengiriman = Pengiriman.query.get_or_404(item_id)
    pengiriman.status = request.form.get("status", pengiriman.status).strip() or pengiriman.status
    db.session.commit()
    return redirect(url_for("kasir_pengiriman"))


@app.route("/gudang/dashboard")
def gudang_dashboard():
    if "user_id" not in session:
        return redirect("/login")

    # Barang masuk bulan ini
    # SQLite tidak mendukung fungsi month()/year() secara default.
    # Gunakan range tanggal untuk ambil data bulan berjalan.
    now = datetime.now()
    start_date = datetime(now.year, now.month, 1)
    if now.month == 12:
        end_date = datetime(now.year + 1, 1, 1)
    else:
        end_date = datetime(now.year, now.month + 1, 1)

    barang_masuk = BarangMasuk.query.filter(
        BarangMasuk.tanggal_masuk >= start_date.date(),
        BarangMasuk.tanggal_masuk < end_date.date(),
    ).all()


    total_masuk = sum(x.jumlah for x in barang_masuk)

    # Total stok tersedia
    total_stok = db.session.query(
        func.sum(Barang.stok)
    ).scalar() or 0

    # Barang keluar bulan ini
    detail = DetailTransaksi.query.all()

    total_keluar = sum(x.qty for x in detail)

    return render_template(
    "gudang/dashboard.html",
    total_masuk=total_masuk,
    total_keluar=total_keluar,
    total_stok=total_stok,
    sekarang=datetime.now()
)

@app.route("/gudang/barang-masuk", methods=["GET", "POST"])
def barang_masuk():
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() not in ["gudang", "staf gudang"]:
        return "Akses Ditolak: Anda bukan Staff Gudang", 403

    barang_list = Barang.query.order_by(Barang.id).all()
    histori = [
        SimpleNamespace(
            id=1,
            barang=SimpleNamespace(nama_barang="Keripik Balado"),
            supplier="PT Sumber Jaya",
            jumlah=50,
            tanggal_masuk="2026-07-01",
            tanggal_expired="2026-12-31",
        )
    ]

    if request.method == "POST":
        barang_id = request.form["barang_id"]
        supplier = request.form["supplier"]
        jumlah = int(request.form["jumlah"])

        # Validasi jumlah
        if jumlah <= 0:
            flash("Jumlah barang harus lebih dari 0.", "error")
            return redirect("/gudang/barang-masuk")

        # Validasi supplier
        if not supplier.strip():
            flash("Supplier tidak boleh kosong.", "error")
            return redirect("/gudang/barang-masuk")

        tanggal_masuk = datetime.strptime(
            request.form["tanggal_masuk"],
            "%Y-%m-%d"
        ).date()

        tanggal_expired = datetime.strptime(
            request.form["tanggal_expired"],
            "%Y-%m-%d"
        ).date()

        # Validasi tanggal
        if tanggal_expired < tanggal_masuk:
            flash("Tanggal kadaluarsa tidak boleh sebelum tanggal masuk.", "error")
            return redirect("/gudang/barang-masuk")

        data = BarangMasuk(
            barang_id=barang_id,
            supplier=supplier,
            jumlah=jumlah,
            tanggal_masuk=tanggal_masuk,
            tanggal_expired=tanggal_expired,
            gudang_id=session["user_id"],
        )

        db.session.add(data)

        barang = Barang.query.get(barang_id)
        if barang:
            barang.stok += jumlah

        db.session.commit()
        flash("Barang masuk berhasil ditambahkan!", "success")
        return redirect("/gudang/barang-masuk")


    barang = Barang.query.all()

    histori = BarangMasuk.query.order_by(
        BarangMasuk.id.desc()
    ).all()

    return render_template(
        "gudang/barang_masuk.html",
        barang=barang,
        histori=histori
    )

@app.route("/gudang/barang-masuk/edit/<int:item_id>", methods=["GET", "POST"])
def edit_barang_masuk(item_id):
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() not in ["gudang", "staf gudang"]:
        return "Akses Ditolak: Anda bukan Staff Gudang", 403

    barang_list = Barang.query.order_by(Barang.id).all()
    data = SimpleNamespace(
        id=item_id,
        barang_id=1,
        supplier="PT Sumber Jaya",
        jumlah=50,
        tanggal_masuk="2026-07-01",
        tanggal_expired="2026-12-31",
    )

    if request.method == "POST":
        return redirect(url_for("barang_masuk"))

    return render_template("gudang/edit_barang_masuk.html", barang_list=barang_list, data=data)


@app.route("/gudang/barang-masuk/hapus/<int:item_id>")
def hapus_barang_masuk(item_id):
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() not in ["gudang", "staf gudang"]:
        return "Akses Ditolak: Anda bukan Staff Gudang", 403

    return redirect(url_for("barang_masuk"))


@app.route("/gudang/laporan-stok")
def laporan_stok():
    if "user_id" not in session:
        return redirect("/login")

    if (session.get("role") or "").lower() not in ["gudang", "staf gudang"]:
        return "Akses Ditolak: Anda bukan Staff Gudang", 403

    keyword = request.args.get("keyword", "").strip()
    query = Barang.query
    if keyword:
        query = query.filter(Barang.nama_barang.ilike(f"%{keyword}%"))

    barang_list = query.order_by(Barang.id).all()
    return render_template("gudang/laporan_stok.html", barang=barang_list, keyword=keyword)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


if __name__ == "__main__":
    app.run(debug=True)